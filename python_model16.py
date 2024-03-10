from docx import Document
from docx.shared import Inches, Pt
from sklearn.preprocessing import LabelEncoder
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, confusion_matrix
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import shutil
import os

# Ensure matplotlib and seaborn settings for visualizations
plt.rcParams.update({'font.size': 12, 'figure.figsize': (8, 6)})

# Load the dataset
dataset_path = 'jobs_in_data.csv'
jobs_data = pd.read_csv(dataset_path)

# Keep a copy of the original job titles for reporting purposes
job_titles_original = jobs_data['job_title'].copy()

# Initialize LabelEncoder
le = LabelEncoder()

# Define columns
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'

# Encode variables
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Prepare features and target
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Split dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Train RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Model evaluation
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)

# Initialize Document for reporting
doc = Document()
doc.add_heading('Data Analysis and Model Performance Report', level=1)
doc.add_paragraph(f'Accuracy of the model: {accuracy:.2f}', style='BodyText')

# Set the output directory and delete if exists
output_dir_path = 'C:\\Users\\baltz\\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\\Επιφάνεια εργασίας\\Model Report & Graphs'
if os.path.exists(output_dir_path):
    shutil.rmtree(output_dir_path)
os.makedirs(output_dir_path)

output_dir = Path(output_dir_path)

# Confusion Matrix with detailed description
cm = confusion_matrix(y_test, y_pred)
fig, ax = plt.subplots()
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues', ax=ax)
ax.set_xlabel('Predicted labels')
ax.set_ylabel('True labels')
ax.set_title('Confusion Matrix with KPIs', pad=20)
plt.tight_layout()

confusion_matrix_description = """
The Confusion Matrix visualizes the classification accuracy of the model, comparing actual vs. predicted values.
Key Performance Indicators (KPIs) include Accuracy (overall correctness), Precision (correct positive predictions among all positive predictions),
and Recall (correct positive predictions among all actual positives). This matrix helps in identifying the model's strengths and weaknesses in classifying each class.
"""
doc.add_paragraph(confusion_matrix_description).runs[0].font.size = Pt(10)

confusion_matrix_path = output_dir / 'confusion_matrix.png'
plt.savefig(confusion_matrix_path)
doc.add_paragraph('Confusion Matrix:', style='Heading2')
doc.add_picture(str(confusion_matrix_path), width=Inches(6))
plt.close()

# Feature Importances with detailed description
features = pd.Series(clf.feature_importances_, index=X.columns)
fig, ax = plt.subplots()
features.nlargest(10).plot(kind='barh', ax=ax)
ax.set_title('Top 10 Feature Importances with KPIs', pad=20)
plt.tight_layout()

feature_importances_description = """
This graph displays the top 10 features that influence the model's predictions the most.
Key Performance Indicator (KPI) is the Feature Contribution, indicating the relative importance of each feature in making accurate predictions.
Features with higher values have a more significant impact on the decision-making process of the model.
"""
doc.add_paragraph(feature_importances_description).runs[0].font.size = Pt(10)

feature_importances_path = output_dir / 'feature_importances.png'
plt.savefig(feature_importances_path)
doc.add_paragraph('Feature Importances:', style='Heading2')
doc.add_picture(str(feature_importances_path), width=Inches(6))
plt.close()

# Add the section for crucial job titles to hire
most_demanded_job_titles = job_titles_original.value_counts().head(5)
doc.add_heading('Crucial Job Titles to Hire', level=1)
doc.add_paragraph('Based on our analysis, the following job titles are in high demand and crucial to hire:')
for job_title, count in most_demanded_job_titles.items():
    doc.add_paragraph(f'{job_title}: {count} listings', style='ListBullet')

# Recommendations for HR manager
doc.add_heading('HR Manager Recommendations', level=1)
recommendations = """
Based on the model's findings, we recommend focusing on recruitment and development in key areas identified as critical by the feature importance analysis...
"""
doc.add_paragraph(recommendations)

# Save the final document
final_doc_path = output_dir / 'Data_Analysis_and_Model_Performance_Report.docx'
doc.save(final_doc_path)
print(f'Report saved to: {final_doc_path}')

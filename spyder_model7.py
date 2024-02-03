# Import necessary libraries
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score, confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import os
import numpy as np

# Define the dataset path
dataset_path = 'jobs_in_data.csv"'  # Update this path to your actual dataset location

# Load the dataset
dataset_path = 'jobs_in_data.csv'  # Update this path
jobs_data = pd.read_csv(dataset_path)

# Preprocessing
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'  # Make sure this column exists

# Encode categorical and target variables
le = LabelEncoder()
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Preparing the features and target variable
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Splitting the dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Training the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Making predictions and evaluating the model
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
print(f'Accuracy: {accuracy}')

# Define the output directory
output_dir = 'C:\\Users\\baltz\\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\\Επιφάνεια εργασίας\\Model Graphs5'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Save and add the confusion matrix plot to the document
confusion_matrix_path = os.path.join(output_dir, 'confusion_matrix.png')
plt.figure(figsize=(10, 7))
cm = confusion_matrix(y_test, y_pred)
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues')
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.savefig(confusion_matrix_path)
plt.close()

# Save and add the feature importance plot to the document
feature_importances_path = os.path.join(output_dir, 'feature_importances.png')
plt.figure(figsize=(10, 7))
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.savefig(feature_importances_path)
plt.close()

# Creating and saving a report document
doc = Document()
doc.add_heading('Model Performance Report', 0)

# Model Summary
doc.add_heading('1. Model Summary', level=1)
doc.add_paragraph(
    f"The RandomForestClassifier model achieved an accuracy of {accuracy:.2f} on the test set. "
    "This section provides an overview of the model's performance metrics, including precision, "
    "recall, f1-score, and a confusion matrix analysis."
)

# Classification Report
doc.add_heading('2. Classification Report', level=1)
report = classification_report(y_test, y_pred, output_dict=True)
# Dynamically generate classification report text
for label, metrics in report.items():
    if label not in ['accuracy', 'macro avg', 'weighted avg']:
        doc.add_paragraph(
            f"Class {label} - Precision: {metrics['precision']:.2f}, "
            f"Recall: {metrics['recall']:.2f}, F1-Score: {metrics['f1-score']:.2f}"
        )

# Confusion Matrix
doc.add_heading('3. Confusion Matrix', level=1)
doc.add_paragraph(
    "The confusion matrix provides insight into the number of correct and incorrect predictions "
    "made by the model, offering a visual representation of performance across the different classes."
)
doc.add_picture(confusion_matrix_path)

# Feature Importances
doc.add_heading('4. Feature Importances', level=1)
doc.add_paragraph(
    "Understanding which features most significantly impact the model's predictions can help in "
    "focusing data collection efforts and potentially improving model performance. The graph below "
    "illustrates the relative importance of the top features considered by the model."
)
doc.add_picture(feature_importances_path)

# In-depth Analysis
doc.add_heading('5. In-depth Analysis', level=1)
doc.add_paragraph(
    "This section would include a more detailed analysis of the model's performance, discussing "
    "specific areas of strength and weakness, interpreting the confusion matrix and feature "
    "importances in the context of the domain, and suggesting potential strategies for improvement. "
    "Given the automated nature of this report, detailed domain-specific insights should be "
    "added by a domain expert."
)

doc_path = os.path.join(output_dir, 'model_report.docx')
doc.save(doc_path)
print(f'Report saved to: {doc_path}')

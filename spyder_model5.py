# Import necessary libraries
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score, confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import os
import numpy as np

# Define the dataset path
dataset_path = 'jobs_in_data.csv"'  # Update this path to your actual dataset location

# Load the dataset
dataset_path = 'jobs_in_data.csv'  # Update this path
jobs_data = pd.read_csv(dataset_path)

# Preprocessing
# Assuming these columns are in your dataset. Adjust as necessary.
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'  # Assuming this is your target variable

# Encode categorical and target variables
le = LabelEncoder()
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Preparing the features and target variable
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Splitting the dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Training the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Making predictions and evaluating the model
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
print(f'Accuracy: {accuracy}')

# Define the output directory
output_dir = 'C:\\Users\\baltz\\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\\Επιφάνεια εργασίας\\Model Graphs2'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Generating and saving the confusion matrix plot
plt.figure(figsize=(10, 7))
cm = confusion_matrix(y_test, y_pred)
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues')
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.savefig(os.path.join(output_dir, 'confusion_matrix.png'))
plt.close()

# Generating and saving the feature importance plot
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
plt.figure(figsize=(10, 7))
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.savefig(os.path.join(output_dir, 'feature_importances.png'))
plt.close()

# Creating and saving a report document
doc = Document()
doc.add_heading('Model Performance Report', 0)
doc.add_paragraph(f'Accuracy: {accuracy}')
doc.add_paragraph('Classification report:')
doc.add_paragraph(classification_report(y_test, y_pred))
doc_path = os.path.join(output_dir, 'model_report.docx')
doc.save(doc_path)
print(f'Report saved to: {doc_path}')

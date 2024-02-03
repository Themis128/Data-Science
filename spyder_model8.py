# Import necessary libraries
from docx import Document
from docx.shared import Inches
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score, confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os
import numpy as np

# Improved visualization settings
plt.rcParams.update({'font.size': 14, 'figure.figsize': (10, 7)})

# Load the dataset
dataset_path = 'jobs_in_data.csv'
jobs_data = pd.read_csv(dataset_path)

# Preprocessing
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'

# Encode categorical and target variables
le = LabelEncoder()
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Preparing the features and target variable
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Splitting the dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Training the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Making predictions and evaluating the model
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)

# Define the output directory
output_dir = 'C:\\Users\\baltz\\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\\Επιφάνεια εργασίας\\Model Graphs5'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Creating and saving a report document
doc = Document()
doc.add_heading('Model Performance Report', 0)

# Include enhanced descriptions and analyses
doc.add_paragraph(f'Accuracy: {accuracy:.2f}\n')
report_text = """
Here we delve into the model's performance, analyzing the accuracy, precision, recall, and F1-score across different classes. 
These metrics collectively offer insights into the model's ability to correctly identify and classify the various job settings from the dataset. 
Accuracy provides a general sense of the model's overall performance, while the precision, recall, and F1-score offer a deeper understanding of its performance on a class-by-class basis.
"""
doc.add_paragraph(report_text)
classification_report_text = classification_report(y_test, y_pred)
doc.add_paragraph('Classification report:\n')
doc.add_paragraph(classification_report_text)

# Generate, save, and add the confusion matrix plot
confusion_matrix_path = os.path.join(output_dir, 'confusion_matrix.png')
plt.figure()
cm = confusion_matrix(y_test, y_pred)
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues')
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.tight_layout()
plt.savefig(confusion_matrix_path, dpi=300)
plt.close()
doc.add_paragraph('\nConfusion Matrix:\n')
doc.add_picture(confusion_matrix_path, width=Inches(6))

# Generate, save, and add the feature importance plot
feature_importances_path = os.path.join(output_dir, 'feature_importances.png')
plt.figure()
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.tight_layout()
plt.savefig(feature_importances_path, dpi=300)
plt.close()

doc.add_paragraph(
"The top 10 most important features determined by the RandomForestClassifier are shown below." 
"This visualization helps in understanding which features contribute most significantly to the model's decision-making process:")

doc.add_picture(feature_importances_path, width=Inches(6))



# Save the Word document
doc_path = os.path.join(output_dir, 'model_report.docx')
doc.save(doc_path)
print(f'Report saved to: {doc_path}')

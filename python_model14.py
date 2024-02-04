from docx import Document
from docx.shared import Inches
from sklearn.preprocessing import LabelEncoder
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, confusion_matrix
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path

# Ensure matplotlib and seaborn settings for visualizations
plt.rcParams.update({'font.size': 12, 'figure.figsize': (8, 6)})

# Load the dataset
dataset_path = 'jobs_in_data.csv'
jobs_data = pd.read_csv(dataset_path)

# Keep a copy of the original job titles for reporting purposes
job_titles_original = jobs_data['job_title'].copy()

# Initialize LabelEncoder
le = LabelEncoder()

# Define columns
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'

# Encode variables
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Prepare features and target
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Split dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Train RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Model evaluation
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)

# Initialize Document for reporting
doc = Document()
doc.add_heading('Data Analysis and Model Performance Report', level=1)
doc.add_paragraph(f'Accuracy of the model: {accuracy:.2f}\n', style='BodyText')

# Confusion Matrix
cm = confusion_matrix(y_test, y_pred)
fig, ax = plt.subplots()
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues', ax=ax)
ax.set_xlabel('Predicted labels')
ax.set_ylabel('True labels')
ax.set_title('Confusion Matrix')
plt.tight_layout()

# Set the output directory
output_dir = Path('C:/Users/baltz/OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο/Επιφάνεια εργασίας/Model Report & Graphs')
output_dir.mkdir(parents=True, exist_ok=True)

# Save and add Confusion Matrix to the document
confusion_matrix_path = output_dir / 'confusion_matrix.png'
plt.savefig(confusion_matrix_path)
doc.add_paragraph('Confusion Matrix:', style='Heading2')
doc.add_picture(str(confusion_matrix_path), width=Inches(6))

plt.close()

# Feature Importances
features = pd.Series(clf.feature_importances_, index=X.columns)
fig, ax = plt.subplots()
features.nlargest(10).plot(kind='barh', ax=ax)
ax.set_title('Top 10 Feature Importances')
plt.tight_layout()

# Save and add Feature Importances to the document
feature_importances_path = output_dir / 'feature_importances.png'
plt.savefig(feature_importances_path)
doc.add_paragraph('Feature Importances:', style='Heading2')
doc.add_picture(str(feature_importances_path), width=Inches(6))

plt.close()

# Find the most demanded job titles (Top 5 for this example)
most_demanded_job_titles = job_titles_original.value_counts().head(5)

# Add a section in the Word document for crucial job titles to hire
doc.add_heading('Crucial Job Titles to Hire', level=1)
doc.add_paragraph('Based on our analysis, the following job titles are in high demand and crucial to hire:')
for job_title, count in most_demanded_job_titles.items():
    doc.add_paragraph(f'{job_title}: {count} listings', style='ListBullet')

# Recommendations for HR manager
doc.add_heading('HR Manager Recommendations', level=1)
recommendations = """
Based on the model's findings, we recommend focusing on recruitment and development in key areas identified as critical by the feature importance analysis...
"""
doc.add_paragraph(recommendations)

# Save the final document
final_doc_path = output_dir / 'Data_Analysis_and_Model_Performance_Report.docx'
doc.save(final_doc_path)
print(f'Report saved to: {final_doc_path}')

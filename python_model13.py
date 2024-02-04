from docx.shared import RGBColor
from docx import Document
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score, confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
from docx.shared import Inches
import os
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Improved visualization settings
plt.rcParams.update({'font.size': 14, 'figure.figsize': (10, 7)})

# Load the dataset
dataset_path = 'jobs_in_data.csv'  # Adjust the path to where you've stored your dataset
jobs_data = pd.read_csv(dataset_path)

# Aggregate the data to count occurrences for each job title, serving as a demand indicator
job_demand = jobs_data['job_title'].value_counts()

# Define the output directory (ensure this path is correct for your environment)
output_dir = Path('C:/Users/baltz/OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο/Επιφάνεια εργασίας/Model Graphs')
if not os.path.exists(output_dir):
    os.makedirs(output_dir)
    
    
# Preprocessing steps
# Encode categorical and target variables
le = LabelEncoder()
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
target_column = 'work_setting'
for column in categorical_columns + [target_column]:
    jobs_data[column] = le.fit_transform(jobs_data[column])

# Preparing the features and target variable
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data[target_column]

# Splitting the dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Training the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Evaluation
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
    

# Documenting the model performance and crucial job titles for hiring
doc = Document()
doc.add_heading('Model Performance and Hiring Focus Report', level=1)

# Assume previous steps for model training and evaluation are here

# Add a section in the Word document for crucial job titles to hire
top_n = 10  # Adjust based on your preference
top_job_titles_demand = job_demand.head(top_n)

doc.add_heading('Crucial Job Titles to Hire', level=1)
doc.add_paragraph(
    "Based on the dataset analysis, the following job titles are identified as crucial for the company to focus its hiring efforts on:"
)

# List the top N job titles with their demand
for job_title, demand in top_job_titles_demand.items():
    doc.add_paragraph(f'{job_title}: {demand} occurrences', style='ListBullet')


# Print accuracy to the screen
print(f'Accuracy: {accuracy:.2f}')

# Add accuracy to the Word document
doc.add_paragraph(f'Accuracy: {accuracy:.2f}\n')

# Enhanced descriptions and analyses
report_text = """
Here we delve into the model's performance, analyzing the accuracy, precision, recall, and F1-score across different classes. 
These metrics collectively offer insights into the model's ability to correctly identify and classify the various job settings from the dataset. 
Accuracy provides a general sense of the model's overall performance, while the precision, recall, and F1-score offer a deeper understanding of its performance on a class-by-class basis.
"""
doc.add_paragraph(report_text)

# Classification report in matrix format
classification_report_text = classification_report(y_test, y_pred, output_dict=True)
classification_matrix = pd.DataFrame(classification_report_text).transpose()

# Define a function to format cell text color
def format_color(cell):
    cell_value = float(cell.text)
    if cell_value > 0.8:  # Change this threshold based on your criteria for "acceptable"
        cell_color = RGBColor(0, 128, 0)  # Green color
    else:
        cell_color = RGBColor(255, 0, 0)  # Red color
    cell.text = f"{cell_value:.2f}"
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = cell_color

# Add classification report to the Word document
doc.add_heading('Classification Report (Matrix Format)', level=2)  # Use 'level' parameter to specify heading level
doc.add_paragraph()  # Add an empty line
table = doc.add_table(rows=1, cols=len(classification_matrix.columns))
table.autofit = False

# Set the column widths
for i in range(len(classification_matrix.columns)):
    table.cell(0, i).width = Inches(1.2)

# Add column headers
for i, column in enumerate(classification_matrix.columns):
    cell = table.cell(0, i)
    cell.text = column
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add classification report values to the table with color formatting
for i, (index, row) in enumerate(classification_matrix.iterrows()):
    table.add_row().cells
    for j, cell_value in enumerate(row):
        cell = table.cell(i + 1, j)
        cell.text = f"{cell_value:.2f}"  # Ensure cell value is formatted as float
        format_color(cell)

# Generate, save, and add the confusion matrix plot
output_dir = Path(output_dir)  # Convert output_dir to a pathlib Path object
confusion_matrix_path = output_dir / 'confusion_matrix.png'

plt.figure()
cm = confusion_matrix(y_test, y_pred)
sns.heatmap(cm, annot=True, fmt="d", cmap='Blues')
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.tight_layout()
plt.savefig(str(confusion_matrix_path), dpi=300)
plt.close()
doc.add_paragraph('Confusion Matrix:', style='Heading2')  # Use 'style' parameter to specify heading style

# Open the confusion matrix image as a stream and add it to the document
with open(confusion_matrix_path, 'rb') as img_file:
    doc.add_picture(img_file, width=Inches(6))

# Feature importance plot
feature_importances_path = output_dir / 'feature_importances.png'
plt.figure()
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.tight_layout()
plt.savefig(str(feature_importances_path), dpi=300)
plt.close()

doc.add_heading('Top 10 Feature Importances', level=2)  # Use 'level' parameter to specify heading level
doc.add_paragraph(
    "The top 10 most important features determined by the RandomForestClassifier are shown below. "
    "This visualization helps in understanding which features contribute most significantly to the model's decision-making process:")

# Open the feature importances image as a stream and add it to the document
with open(feature_importances_path, 'rb') as img_file:
    doc.add_picture(img_file, width=Inches(6))
    
# Recommendations for HR manager based on the analysis
doc.add_heading('HR Manager Recommendations', level=2)
recommendations = """
Based on the model's findings, we recommend focusing on recruitment and development in key areas identified as critical by the feature importance analysis. For instance, if 'experience_level' and 'job_title' are among the top features, enhancing training programs for existing employees and tailoring recruitment to attract experienced professionals in pivotal roles could be beneficial. Additionally, addressing any significant salary discrepancies to remain competitive and retain talent is advisable. Investing in employee development and reconsidering salary structures could significantly enhance workforce capabilities and satisfaction.
"""
doc.add_paragraph(recommendations)
    
# Strategic Recommendations Section
doc.add_heading('Strategic Recommendations for HR Management', level=1)

# Analyze the confusion matrix and feature importances to derive recommendations
recommendations_text = """
Based on the model's performance and the importance of specific features, it is recommended that HR focuses on enhancing recruitment and development in areas identified as critical, such as positions with high feature importance scores. Adjusting salary scales based on the competitive analysis provided by the salary features could help in attracting top talent. Additionally, areas where the confusion matrix shows high misclassification rates could benefit from targeted job role clarifications and training, to ensure clear role definitions and expectations.
"""

doc.add_paragraph(recommendations_text)

# Ensure only one Document object is initialized
doc = Document()
doc.add_heading('Model Performance and Hiring Focus Report', level=1)

# Model performance and evaluation sections here

# Documenting the model performance
doc = Document()
doc.add_heading('Model Performance Report', level=1)  # Use 'level' parameter to specify heading level
# (These sections have been appropriately adjusted in your script)

# Add crucial job titles to hire section
top_n = 10  # Adjust based on your preference
top_job_titles_demand = job_demand.head(top_n)

doc.add_heading('Crucial Job Titles to Hire', level=2)
doc.add_paragraph(
    "Based on the dataset analysis, the following job titles are identified as crucial for the company to focus its hiring efforts on:"
)

for job_title, demand in top_job_titles_demand.items():
    doc.add_paragraph(f'{job_title}: {demand} occurrences', style='ListBullet')

# Continue with additional analysis and recommendations as per your script
# Print accuracy to the screen
print(f'Accuracy: {accuracy:.2f}')

# Save the final document with recommendations and crucial job titles
final_doc_path = output_dir / 'HR_Strategic_Recommendations_Report.docx'
doc.save(final_doc_path)
print(f'Report with strategic recommendations saved to: {final_doc_path}')
# Import necessary libraries
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score, confusion_matrix, roc_curve, auc, precision_recall_curve, average_precision_score
from sklearn.preprocessing import label_binarize
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import os

# Load the dataset
dataset_path = 'jobs_in_data.csv'  # Update this path
jobs_data = pd.read_csv(dataset_path)

# Preprocessing
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence', 'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']
jobs_data = jobs_data[categorical_columns + numerical_columns + ['work_setting']]  # Assuming 'work_setting' is your target variable

# Encode categorical and target variables
le = LabelEncoder()
for column in categorical_columns:
    jobs_data[column] = le.fit_transform(jobs_data[column])

jobs_data['work_setting_encoded'] = le.fit_transform(jobs_data['work_setting'])
target_names = le.classes_

# Splitting dataset
X = jobs_data[categorical_columns + numerical_columns]
y = jobs_data['work_setting_encoded']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Model training
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Predictions
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
print(f'Accuracy: {accuracy:.2f}')

# Generating classification report
report = classification_report(y_test, y_pred, target_names=target_names)
print(report)

# Plotting Confusion Matrix
cm = confusion_matrix(y_test, y_pred)
plt.figure(figsize=(10,7))
sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", xticklabels=target_names, yticklabels=target_names)
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.show()

# Feature Importances
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
plt.figure(figsize=(10,7))
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.show()

# Checking and creating the output directory
output_dir = r'C:\Users\baltz\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\Επιφάνεια εργασίας\Model_Graphs2'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Saving the model report to a Word document
doc = Document()
doc.add_heading('Model Report for Predicting Work Settings', 0)
doc.add_paragraph(f'Accuracy on test set: {accuracy:.2f}')
doc.add_paragraph('Classification Report')
doc.add_paragraph(report)
doc_path = os.path.join(output_dir, 'model_report.docx')
doc.save(doc_path)
print(f'Model report saved to: {doc_path}')

# Note: For ROC and Precision-Recall Curve, additional preprocessing and plotting are required, especially for multi-class scenarios.

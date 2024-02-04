# Import necessary libraries
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
import os
import numpy as np

# Define the dataset path
dataset_path = 'jobs_in_data.csv"'  # Update this path to your actual dataset location

# Load the dataset
dataset_path = 'jobs_in_data.csv'  # Update this path
jobs_data = pd.read_csv(dataset_path)

# Ensure 'work_setting' column exists and encode it if not already done
if 'work_setting_encoded' not in jobs_data.columns:
    if 'work_setting' in jobs_data.columns:
        le = LabelEncoder()
        jobs_data['work_setting_encoded'] = le.fit_transform(jobs_data['work_setting'])
    else:
        raise ValueError("Column 'work_setting' not found in the dataset.")

# Preprocess and encode other categorical columns
categorical_columns = jobs_data.select_dtypes(include=['object']).columns.tolist()
categorical_columns.remove('work_setting')  # Assuming 'work_setting' is not needed as a feature
le = LabelEncoder()
for column in categorical_columns:
    jobs_data[column] = le.fit_transform(jobs_data[column].astype(str))

# Specify features and target
X = jobs_data.drop(['work_setting_encoded', 'work_setting'], axis=1, errors='ignore')
y = jobs_data['work_setting_encoded']

# Split the dataset
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Train the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Predict and evaluate the model
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
print(f'Accuracy: {accuracy}')

# Output directory
output_dir = r'C:\Users\baltz\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\Επιφάνεια εργασίας\Model Graphs2'
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Confusion Matrix
cm = confusion_matrix(y_test, y_pred)
plt.figure(figsize=(10, 7))
sns.heatmap(cm, annot=True, fmt="d")
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title('Confusion Matrix')
plt.savefig(os.path.join(output_dir, 'confusion_matrix.png'))
plt.close()

# Feature Importances
feature_importances = pd.Series(clf.feature_importances_, index=X.columns)
plt.figure(figsize=(10, 7))
feature_importances.nlargest(10).plot(kind='barh')
plt.title('Feature Importances')
plt.savefig(os.path.join(output_dir, 'feature_importances.png'))
plt.close()

# Generate a Word document report
doc = Document()
doc.add_heading('Model Training Report', 0)
doc.add_paragraph(f'Model accuracy: {accuracy:.2f}')
doc_path = os.path.join(output_dir, 'model_report.docx')
doc.save(doc_path)

print(f'Report and plots saved to {output_dir}')

# Import necessary libraries
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, accuracy_score
from docx import Document
import os

# Define the output paths
output_path = r'C:\Users\baltz\OneDrive - Ελληνικό Ανοικτό Πανεπιστήμιο\Επιφάνεια εργασίας\Model_Graphs1'
doc_path = os.path.join(output_path, 'model_report_for_remote_work_opportunities.docx')

# Ensure the output directory exists
os.makedirs(output_path, exist_ok=True)

# Load and preprocess dataset (Make sure to replace 'jobs_in_data.csv' with the actual path to your dataset)
dataset_path = 'jobs_in_data.csv'  # Update this path
jobs_data = pd.read_csv(dataset_path)

# Assuming 'job_title', 'job_category', etc., are categorical columns in your dataset
categorical_columns = ['job_title', 'job_category', 'salary_currency', 'employee_residence',
                       'experience_level', 'employment_type', 'company_location', 'company_size']
numerical_columns = ['work_year', 'salary', 'salary_in_usd']

# Encode categorical variables
le = LabelEncoder()
jobs_data_encoded = jobs_data.copy()
for column in categorical_columns:
    jobs_data_encoded[column] = le.fit_transform(jobs_data[column])

# Assuming 'work_setting' is a target variable, ensure it's encoded
if 'work_setting' in jobs_data.columns:
    jobs_data_encoded['work_setting_encoded'] = le.fit_transform(jobs_data['work_setting'])
    target_names = le.classes_
else:
    raise ValueError("Column 'work_setting' not found in dataset. Please check the correct target variable.")

# Prepare features and target variable
X = jobs_data_encoded[categorical_columns + numerical_columns]
y = jobs_data_encoded['work_setting_encoded']

# Split the dataset into training and testing sets
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# Train the RandomForestClassifier
clf = RandomForestClassifier(random_state=42)
clf.fit(X_train, y_train)

# Make predictions and evaluate the model
y_pred = clf.predict(X_test)
accuracy = accuracy_score(y_test, y_pred)
classification_report_str = classification_report(y_test, y_pred, target_names=target_names)

# Create and save a Word document with the model report
doc = Document()
doc.add_heading('Model Report for Predicting Remote Work Opportunities', 0)
doc.add_heading('Model Performance', level=1)
doc.add_paragraph(f'Accuracy on test set: {accuracy:.2f}\n\n')
doc.add_heading('Classification Report', level=1)
doc.add_paragraph(classification_report_str)

# Save the document
doc.save(doc_path)

print(f'Model report saved to: {doc_path}')
c